#!/usr/bin/env python3
"""
Build the J Kiss LLC Independent Contractor Agreement.

This script is the single source for BOTH published artifacts:

    J-Kiss-LLC-Independent-Contractor-Agreement-v1.0.docx   (editable master)
    J-Kiss-LLC-Independent-Contractor-Agreement-v1.0.pdf    (rendered from the DOCX)

The PDF is produced by converting the generated DOCX with LibreOffice, so the two
files cannot drift apart in wording. Edit the CONTENT block below and re-run:

    python3 docs/legal/build-contractor-agreement.py

Requires: python-docx, and LibreOffice (`soffice`) on PATH for the PDF step.

STATUS: approved by the owner for use in Operion on August 21, 2026.
"""

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── Identity (mirrors app/lib/company.ts) ────────────────────────────────────
COMPANY_NAME = "J Kiss LLC"
COMPANY_ADDRESS = "8055 Windrose Ave, Plano, TX 75024"
COMPANY_PHONE = "(817) 909-4312"
COMPANY_EMAIL = "info@jkissllc.com"
COMPANY_SITE = "jkissllc.com"
COMPANY_DOT = "US DOT 3484556 / MC 01155352"

VERSION = "1.0"
REVISED = "August 20, 2026"
DOC_TITLE = "Independent Contractor Agreement"

INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
RULE = "999999"

OUT_DIR = Path(__file__).resolve().parent
BASENAME = "J-Kiss-LLC-Independent-Contractor-Agreement-v1.0"


# ── Low-level docx helpers ───────────────────────────────────────────────────
def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        spec = kwargs.get(edge)
        if not spec:
            continue
        el = OxmlElement(f"w:{edge}")
        for key, value in spec.items():
            el.set(qn(f"w:{key}"), str(value))
        borders.append(el)
    tc_pr.append(borders)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), fill)
    tc_pr.append(el)


def cell_margins(cell, left=110, right=110, top=60, bottom=60):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, width in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(width))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tc_pr.append(mar)


def set_row_height(row, inches):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:trHeight")
    el.set(qn("w:val"), str(int(inches * 1440)))
    el.set(qn("w:hRule"), "atLeast")
    tr_pr.append(el)


def keep_with_next(paragraph):
    """Stop a heading from being orphaned at the foot of a page."""
    paragraph.paragraph_format.keep_with_next = True


def keep_together(paragraph):
    paragraph.paragraph_format.keep_together = True


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), RULE)
    borders.append(bottom)
    p_pr.append(borders)
    return p


# ── Document construction ────────────────────────────────────────────────────
def build_document() -> Document:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)

    build_footer(section)
    build_cover(doc)
    build_body(doc)
    build_signatures(doc)
    build_exhibit_a(doc)
    return doc


def build_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"{COMPANY_NAME} · {DOC_TITLE} · v{VERSION} ({REVISED})    |    Page "
    )
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    add_field(p, " PAGE ")
    run = p.add_run(" of ")
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    add_field(p, " NUMPAGES ")
    for r in p.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = MUTED


def build_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(COMPANY_NAME.upper())
    run.bold = True
    run.font.size = Pt(20)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{COMPANY_ADDRESS}  ·  {COMPANY_PHONE}\n{COMPANY_EMAIL}  ·  {COMPANY_SITE}  ·  {COMPANY_DOT}")
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INDEPENDENT CONTRACTOR AGREEMENT")
    run.bold = True
    run.font.size = Pt(15)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Version {VERSION}  ·  Revised {REVISED}")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(12)

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    keep_with_next(p)
    return p


def body(doc, text, indent=True, together=False):
    p = doc.add_paragraph(text)
    if indent:
        p.paragraph_format.left_indent = Inches(0.18)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Widow/orphan control: never strand one or two lines of a clause on its own page.
    p.paragraph_format.widow_control = True
    if together:
        keep_together(p)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.42)
        p.paragraph_format.space_after = Pt(3)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def build_body(doc):
    n = [0]

    def sec(title):
        n[0] += 1
        h1(doc, f"{n[0]}.  {title}")

    # 1
    sec("Parties and Effective Date")
    body(doc,
         f"This Independent Contractor Agreement (this “Agreement”) is entered into by and between {COMPANY_NAME}, "
         f"a Texas limited liability company with its principal place of business at {COMPANY_ADDRESS} (the “Company”), "
         "and the individual or entity identified below (the “Contractor”). This Agreement takes effect on the "
         "Effective Date stated in Exhibit A, or on the date of the Contractor’s signature if no Effective Date is stated.")
    contractor_identity_block(doc)

    # 2
    sec("Engagement and Scope of Services")
    body(doc,
         "The Company may from time to time offer the Contractor discrete assignments consisting of driving, moving, "
         "loading, hauling, junk removal, delivery, or related services (the “Services”). The specific Services, the "
         "compensation for them, and any special requirements are set out in the applicable work order or assignment "
         "record described in Section 8. This Agreement governs the terms of the relationship; it does not itself "
         "commit either party to any particular assignment.")

    # 3
    sec("Independent Contractor Relationship")
    body(doc,
         "The parties intend that the Contractor perform the Services as an independent contractor and not as an "
         "employee, agent, partner, or joint venturer of the Company. The Contractor is engaged in an independent "
         "trade, occupation, or business.")
    body(doc,
         "The parties acknowledge and agree that this Agreement, by itself, does not determine the Contractor’s legal "
         "classification. Worker classification is determined by the actual working relationship between the parties "
         "and by applicable federal and Texas law as interpreted and applied by the relevant authorities — including "
         "the Internal Revenue Service, the United States Department of Labor, and the Texas Workforce Commission — "
         "as those standards exist from time to time. A label, a written agreement, or the issuance of an IRS Form "
         "1099 does not, by itself, establish independent contractor status. Each party is responsible for conducting "
         "itself consistently with the independent relationship described in this Agreement. Either party may seek a "
         "determination of status from the appropriate authority, including by filing IRS Form SS-8.")

    # 4
    sec("Control Over Means and Methods")
    body(doc,
         "The Contractor controls the manner, means, and methods by which the Services are performed, including the "
         "sequence of tasks, the techniques used, and the decision whether to engage qualified assistants or "
         "substitutes at the Contractor’s own expense, subject to the Contractor remaining responsible for the "
         "Services and for any person the Contractor engages.")
    body(doc, "The Company may specify only the following, which define the result to be achieved rather than the method:")
    bullets(doc, [
        "customer requirements, service specifications, and the result to be delivered;",
        "site access windows, appointment times, and delivery or completion deadlines agreed with the customer;",
        "safety rules and site rules applicable at customer or third-party premises;",
        "requirements imposed by law, regulation, permit, insurance, or a customer contract; and",
        "documentation the Company must retain to satisfy its own legal, tax, insurance, or customer obligations.",
    ])

    # 5
    sec("Assignments; No Guarantee of Work")
    body(doc,
         "The Contractor is free to accept or decline any assignment offered, for any reason, without penalty and "
         "without affecting the Contractor’s eligibility for future offers. The Company does not guarantee any "
         "assignments, hours, routes, schedule, volume of work, or minimum compensation. Nothing in this Agreement "
         "obligates the Company to offer work or the Contractor to accept it. Once the Contractor accepts an "
         "assignment, the Contractor is responsible for performing it or for giving the Company reasonable advance "
         "notice so the Company can make alternative arrangements.")

    # 6
    sec("Services for Others")
    body(doc,
         "The Contractor may perform services for other businesses, including competitors of the Company, and may "
         "advertise and hold out the Contractor’s services to the public. The Contractor’s only limits are the "
         "confidentiality obligations in Section 17 and the obligation to perform assignments the Contractor has "
         "already accepted.")

    # 7
    sec("No Authority to Bind")
    body(doc,
         "The Contractor has no authority to enter into contracts, incur obligations, make representations, settle "
         "claims, or otherwise bind the Company, and shall not hold out any such authority. The Contractor shall not "
         "represent to any person that the Contractor is an employee of the Company.")

    # 8
    sec("Work Orders; Assignment Records")
    body(doc,
         "Each accepted assignment is described in a work order — which may take the form of the assignment record "
         "generated in the Company’s Operion platform, a written work order, or Exhibit A to this Agreement. The "
         "applicable work order controls the description of Services, the compensation basis and rate, any approved "
         "expenses, and any special customer, route, equipment, or safety requirements for that assignment.")
    body(doc,
         "Compensation is determined by the accepted work order for each assignment. This Agreement does not set a "
         "universal rate. Where a work order conflicts with this Agreement on assignment-specific commercial terms, "
         "the work order controls for that assignment; this Agreement controls on all other terms.")

    # 9
    sec("Payment Schedule")
    body(doc,
         "The Company issues contractor payments on a weekly cycle, on Fridays. An assignment becomes payable on the "
         "Friday on or following the close of the pay period in which the assignment was completed and the required "
         "documentation was submitted. Assignments completed, or documentation submitted, after a cycle has closed "
         "are carried to the next Friday cycle. The Company may reasonably delay payment for a specific assignment "
         "pending resolution of a documented customer dispute, a damage claim, or missing required documentation, and "
         "will tell the Contractor the reason for the delay.")
    body(doc,
         "The Company issues a written pay statement for each payment showing the assignments covered, the amounts, "
         "and any agreed deductions. The Contractor should review each statement and raise any question promptly.")

    # 10
    sec("Taxes; Form W-9")
    body(doc,
         "The Contractor is solely responsible for all federal, state, and local taxes arising from the compensation "
         "paid under this Agreement, including income tax, self-employment tax, and any applicable business or "
         "franchise taxes, and for making any required estimated tax payments. The Company will not withhold income "
         "tax, Social Security, or Medicare from payments to the Contractor, and will not pay the employer portion of "
         "any employment tax, except where withholding or reporting is required by law (for example, backup "
         "withholding).")
    body(doc,
         "The Contractor shall provide a complete and accurate IRS Form W-9 before the first payment and shall provide "
         "an updated Form W-9 promptly upon any change to the Contractor’s legal name, entity type, address, or "
         "taxpayer identification number. The Company will report payments on IRS Form 1099-NEC where required. The "
         "Company retains only the last four digits of the Contractor’s taxpayer identification number outside of the "
         "Contractor’s secured Form W-9 record.")

    # 11
    sec("No Employee Benefits")
    body(doc,
         "The Contractor is not eligible for and waives any claim to employee benefits of the Company, including "
         "health, dental, vision, life, or disability insurance; retirement or pension contributions; paid or unpaid "
         "leave; holiday, vacation, or sick pay; severance; workers’ compensation coverage; unemployment "
         "compensation; or overtime pay, except where any such benefit or protection is required by applicable law "
         "notwithstanding the parties’ characterization of the relationship. The Company does not carry workers’ "
         "compensation insurance covering the Contractor. Texas law does not require most private employers to carry "
         "workers’ compensation insurance, and the Contractor is encouraged to obtain the Contractor’s own coverage "
         "for occupational injury.")

    # 12
    sec("Expenses and Reimbursement")
    body(doc,
         "Except as expressly stated in an accepted work order, the Contractor bears the Contractor’s own costs of "
         "doing business, including vehicle acquisition, fuel, maintenance, tolls, insurance premiums, licensing, "
         "mobile service, personal protective equipment, and general tools. The Company reimburses only expenses "
         "that are identified as reimbursable in the applicable work order or approved in writing in advance, and "
         "that are supported by an itemized receipt submitted within thirty (30) days of the expense. Approved "
         "reimbursements are paid on the normal weekly cycle and are not compensation for Services.")

    # 13
    sec("Equipment, Vehicle, Licensing, and Regulatory Compliance")
    body(doc,
         "Where an assignment requires a vehicle, equipment, or a credential, the Contractor is responsible for "
         "providing and maintaining it in safe, lawful, and working condition, unless the work order states that the "
         "Company will supply it. The Contractor shall:")
    bullets(doc, [
        "hold and maintain every licence, endorsement, permit, medical certification, and registration required for "
        "the Services the Contractor performs, and keep each current and unexpired;",
        "comply with all applicable motor-carrier, transportation, environmental, disposal, and occupational safety "
        "laws and regulations applicable to the Services;",
        "notify the Company promptly, and before performing any further driving assignment, if a required licence or "
        "certification is suspended, revoked, restricted, or expires; and",
        "not perform any assignment while impaired, uninsured, unlicensed, or otherwise not lawfully able to perform it.",
    ])
    body(doc,
         "Company-supplied equipment remains Company property, must be used only for Company assignments, and must be "
         "returned in accordance with Section 25.")

    # 14
    sec("Insurance")
    body(doc,
         "The Contractor shall maintain, at the Contractor’s own expense, the insurance coverage appropriate to the "
         "Services the Contractor performs and required by the applicable work order. Where the Contractor operates a "
         "vehicle in connection with the Services, the Contractor shall maintain at least the automobile liability "
         "coverage required by Texas law and any higher limit stated in the work order, together with coverage "
         "appropriate to the commercial use of the vehicle. The Company may require the Contractor to furnish a "
         "certificate of insurance and to give notice of cancellation or material change. Insurance requirements vary "
         "by role; a contractor who performs only non-driving helper services is not required to carry commercial "
         "automobile coverage unless the work order says so.")

    # 15
    sec("Safety, Customer Property, and Incident Reporting")
    body(doc,
         "The Contractor shall perform the Services safely and shall treat customer premises and property with care. "
         "The Contractor shall report to the Company, as soon as reasonably possible and in any event within "
         "twenty-four (24) hours, any accident, injury, property damage, loss, theft, citation, customer complaint, "
         "or other incident arising from or occurring during an assignment. The Contractor shall cooperate reasonably "
         "with the Company’s and any insurer’s investigation of a reported incident. Reporting an incident in good "
         "faith is not, by itself, a basis for ending the relationship.")

    # 16
    sec("Lawful and Professional Conduct")
    body(doc,
         "The Contractor shall comply with applicable law while performing the Services, shall not perform the "
         "Services while under the influence of alcohol or any substance that impairs safe performance, and shall "
         "conduct itself professionally and courteously with customers, the public, and other contractors. The "
         "Contractor shall not engage in unlawful discrimination or harassment in connection with the Services.")

    # 17
    sec("Confidentiality")
    body(doc,
         "“Confidential Information” means non-public information the Contractor receives in connection with the "
         "Services, including customer names, addresses, contact details, and access codes; pricing and rate "
         "information; route and scheduling information; business methods; and information about other contractors. "
         "Confidential Information does not include information that is or becomes public through no fault of the "
         "Contractor, that the Contractor already lawfully possessed, or that the Contractor develops independently.")
    body(doc,
         "The Contractor shall use Confidential Information only to perform accepted assignments, shall not disclose "
         "it to any third party, and shall not use it to solicit the Company’s customers for the Contractor’s own "
         "account or for another business. This obligation continues for two (2) years after the last assignment, "
         "except that information constituting a trade secret remains protected for as long as it qualifies as a "
         "trade secret under applicable law. Nothing in this Agreement prevents the Contractor from disclosing "
         "information where required by law or legal process, from reporting a suspected violation of law to a "
         "government agency, or from discussing the Contractor’s own compensation or working conditions.")

    # 18
    sec("Data Privacy and Use of the Operion Platform")
    body(doc,
         "The Company operates a business platform called Operion, which the Contractor may use to receive assignment "
         "offers, submit documentation, and view pay statements. The Contractor shall keep the Contractor’s access "
         "credentials confidential, shall not share access with any other person, and shall notify the Company "
         "promptly if the Contractor believes access has been compromised.")
    body(doc,
         "The Company collects and stores the Contractor’s onboarding, identity, tax, insurance, assignment, and "
         "payment records to operate its business and to meet its legal, tax, insurance, and customer obligations. "
         "Sensitive documents such as the Contractor’s Form W-9, identity documents, and executed agreement are "
         "stored in encrypted form and are accessible only to authorized Company administrators and to the "
         "Contractor. The Company retains records for the periods required by applicable law and its records policy, "
         "and preserves records subject to a legal hold. The Contractor shall handle any customer personal "
         "information encountered during an assignment only as needed to perform that assignment.")

    # 19
    sec("Records and Document Accuracy")
    body(doc,
         "The Contractor shall provide accurate, complete, and current information and documents to the Company, "
         "including the Contractor’s legal name, mailing address, taxpayer identification, licences, certifications, "
         "and insurance. Submitting a document that the Contractor knows to be false, altered, expired, or belonging "
         "to another person is a material breach of this Agreement. The Contractor shall notify the Company promptly "
         "when any submitted document expires or ceases to be accurate.")

    # 20
    sec("Onboarding, Verification, and Eligibility for Assignments")
    body(doc,
         "Acceptance of the Contractor’s application does not by itself make the Contractor eligible for assignments. "
         "Before the Contractor may be offered or accept assignments, the Contractor must complete onboarding and the "
         "Company must verify it. Onboarding consists of:")
    bullets(doc, [
        "downloading the specific version of this Agreement issued to the Contractor with the Contractor’s onboarding "
        "request, signing it, and returning the executed copy through the secure onboarding page;",
        "submitting a complete and accurate IRS Form W-9; and",
        "submitting the role-specific documents identified during onboarding, which may include a valid driver’s "
        "licence, evidence of insurance, and a badge photograph.",
    ])
    body(doc,
         "The version of this Agreement issued with an onboarding request is fixed for that request. If the Company "
         "later revises its agreement template, the revision does not change the version the Contractor was asked to "
         "sign; a revised version applies only to a new onboarding request. A Company administrator reviews the "
         "submitted documents, and the Contractor becomes eligible for assignments only after that review is "
         "completed and recorded. Submitting documents does not by itself create eligibility.")

    # 21
    sec("No Non-Competition Restriction")
    body(doc,
         "This Agreement contains no non-competition covenant. The Contractor is free during and after the term of "
         "this Agreement to perform services of any kind for any person, including competitors of the Company, and to "
         "operate the Contractor’s own business. The only continuing restrictions are the confidentiality and "
         "non-use obligations in Section 17, which are limited to the Company’s Confidential Information and do not "
         "restrict the Contractor’s ability to earn a living.")

    # 22
    sec("Responsibility and Indemnification")
    body(doc,
         "Each party shall be responsible for, and shall indemnify and hold harmless the other party from, third-party "
         "claims, damages, and reasonable defence costs to the extent caused by that party’s own negligence, wilful "
         "misconduct, or breach of this Agreement. Neither party is required to indemnify the other for any portion "
         "of a claim caused by the other party’s own negligence or wilful misconduct.")
    body(doc,
         "Neither party is liable to the other for indirect, incidental, consequential, special, exemplary, or "
         "punitive damages, or for lost profits, arising out of this Agreement. These limits do not apply to a party’s "
         "indemnification obligations for third-party claims, to a breach of Section 17, or to any liability that "
         "cannot be limited under applicable law. The indemnified party shall give prompt notice of any claim and "
         "shall reasonably cooperate in its defence.")

    # 23
    sec("Term and Termination")
    body(doc,
         "This Agreement begins on the Effective Date and continues until terminated. Either party may terminate this "
         "Agreement at any time, with or without cause, by giving written notice to the other. Either party may "
         "terminate immediately for a material breach, for the loss or expiry of a licence, certification, or "
         "insurance required for the Services, or for conduct that presents a safety or legal risk.")

    # 24
    sec("Effect of Termination")
    body(doc,
         "Upon termination, the Company will stop offering assignments and will end the Contractor’s access to the "
         "Operion platform. Unless the termination is for a safety or legal risk that makes completion inappropriate, "
         "the Contractor shall either complete any assignment already accepted and in progress or coordinate an "
         "orderly handoff to the Company. The Company shall pay the Contractor for all Services properly performed "
         "and documented before termination, together with approved reimbursable expenses, on the normal weekly "
         "cycle. Termination does not affect either party’s right to payment or reimbursement already earned.")
    body(doc,
         "The Company retains the Contractor’s records after termination for the periods required by applicable law "
         "and its records policy, including tax and insurance records. If the parties later resume the relationship, "
         "the Contractor must complete onboarding and verification again as set out in Section 20; resuming does not "
         "bypass any document or verification requirement.")
    body(doc,
         "Sections 7, 10, 11, 15, 17, 18, 19, 21, 22, 24, 25, 26, and 27 through 33 survive termination, together "
         "with any other provision that by its nature should survive.")

    # 25
    sec("Return of Company Property")
    body(doc,
         "Within ten (10) business days after termination, or sooner on the Company’s request, the Contractor shall "
         "return all Company property in the Contractor’s possession, including equipment, keys, access cards, "
         "uniforms, badges, customer lists, and documents containing Confidential Information, and shall delete or "
         "destroy any remaining copies of Confidential Information in the Contractor’s control.")

    # 26
    sec("Governing Law and Venue")
    body(doc,
         "This Agreement is governed by the laws of the State of Texas, without regard to its conflict-of-laws rules. "
         "The parties agree that the exclusive venue for any action arising out of or relating to this Agreement "
         "shall be the state or federal courts located in or serving Collin County, Texas, and each party consents to "
         "the personal jurisdiction of those courts. The parties may agree in writing to mediate a dispute before "
         "filing suit.")

    # 27
    sec("Notices")
    body(doc,
         f"Notices under this Agreement must be in writing and are effective when delivered by hand, by nationally "
         f"recognized overnight courier, by certified mail return receipt requested, or by email to the address the "
         f"party has most recently provided. Notices to the Company go to {COMPANY_NAME}, {COMPANY_ADDRESS}, or "
         f"{COMPANY_EMAIL}. Notices to the Contractor go to the mailing address or email address in Exhibit A. Each "
         "party shall keep its notice information current.")

    # 28
    sec("Assignment")
    body(doc,
         "The Contractor may not assign this Agreement or delegate the Contractor’s obligations under it without the "
         "Company’s prior written consent, except that the Contractor may engage qualified assistants as permitted by "
         "Section 4. The Company may assign this Agreement to a successor in connection with a merger, "
         "reorganization, or sale of substantially all of its assets. This Agreement binds and benefits the parties "
         "and their permitted successors and assigns.")

    # 29
    sec("Waiver")
    body(doc,
         "A party’s failure or delay in enforcing any provision of this Agreement is not a waiver of that provision "
         "or of any other provision. A waiver is effective only if it is in writing and signed by the waiving party, "
         "and applies only to the specific instance stated.")

    # 30
    sec("Severability")
    body(doc,
         "If any provision of this Agreement is held invalid, illegal, or unenforceable, that provision shall be "
         "modified to the minimum extent necessary to make it enforceable, or if it cannot be modified, severed, and "
         "the remaining provisions shall remain in full force and effect.")

    # 31
    sec("Entire Agreement")
    body(doc,
         "This Agreement, together with any accepted work order and Exhibit A, is the entire agreement between the "
         "parties regarding its subject matter and supersedes all prior or contemporaneous discussions, "
         "representations, and agreements, whether oral or written, on that subject.")

    # 32
    sec("Amendments")
    body(doc,
         "This Agreement may be amended only by a writing signed by both parties. A revised version of the Company’s "
         "agreement template does not amend an agreement already executed by the Contractor; a revised version "
         "applies only if the Contractor executes it.")

    # 33
    sec("Counterparts and Electronic Signatures")
    body(doc,
         "This Agreement may be executed in counterparts, each of which is an original and all of which together form "
         "one agreement. The parties agree that electronic signatures, typed signatures adopted by the signer, and "
         "signatures delivered by electronic image or through the Company’s secure onboarding page are valid and "
         "binding to the same extent as original handwritten signatures, consistent with the Texas Uniform Electronic "
         "Transactions Act and the federal E-SIGN Act.", together=True)


def contractor_identity_block(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    rows = [
        ("Contractor legal name", ""),
        ("Business name (if any)", ""),
        ("Role", "☐ Driver    ☐ Helper    ☐ Other: ______________________"),
        ("Mailing address", ""),
        ("Email", ""),
        ("Phone", ""),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(1.9)
    table.columns[1].width = Inches(4.6)
    for i, (label, value) in enumerate(rows):
        left, right = table.rows[i].cells
        left.width = Inches(1.9)
        right.width = Inches(4.6)
        shade(left, "F4F4F6")
        for cell, text, bold in ((left, label, True), (right, value, False)):
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(9.5)
        for cell in (left, right):
            cell_margins(cell)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": "CCCCCC"},
                bottom={"val": "single", "sz": 4, "color": "CCCCCC"},
                left={"val": "single", "sz": 4, "color": "CCCCCC"},
                right={"val": "single", "sz": 4, "color": "CCCCCC"},
            )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def signature_block(doc, heading, lines):
    p = doc.add_paragraph()
    run = p.add_run(heading)
    run.bold = True
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    keep_with_next(p)
    keep_together(p)
    for label in lines:
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(10)
        sp.paragraph_format.left_indent = Inches(0.12)
        keep_together(sp)
        keep_with_next(sp)
        run = sp.add_run(label)
        run.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_signatures(doc):
    # Keep the whole signature area on one page.
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    p = doc.add_paragraph()
    run = p.add_run("SIGNATURES")
    run.bold = True
    run.font.size = Pt(13)
    p.paragraph_format.space_after = Pt(4)
    keep_with_next(p)
    horizontal_rule(doc)

    p = doc.add_paragraph()
    run = p.add_run(
        "By signing below, each party states that it has read and understood this Agreement and agrees to be bound "
        "by it. The Contractor states that the information the Contractor has provided is accurate and complete."
    )
    run.font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(16)
    keep_with_next(p)

    signature_block(doc, f"{COMPANY_NAME.upper()}", [
        "Signature:  ______________________________________________",
        "Printed name:  ___________________________________________",
        "Title:  __________________________________________________",
        "Date:  ___________________________________________________",
    ])

    signature_block(doc, "CONTRACTOR", [
        "Signature:  ______________________________________________",
        "Printed name:  ___________________________________________",
        "Business name (if signing on behalf of an entity):  _______________________",
        "Title (if applicable):  ___________________________________",
        "Date:  ___________________________________________________",
    ])

    p = doc.add_paragraph()
    run = p.add_run(
        "A typed name adopted by the signer, or a signature submitted through the Company’s secure onboarding page, "
        "is accepted as a valid signature under Section 33."
    )
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED
    keep_together(p)


def build_exhibit_a(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    p = doc.add_paragraph()
    run = p.add_run("EXHIBIT A — WORK ORDER / ASSIGNMENT ADDENDUM")
    run.bold = True
    run.font.size = Pt(13)
    p.paragraph_format.space_after = Pt(2)
    keep_with_next(p)

    p = doc.add_paragraph()
    run = p.add_run(
        "This Exhibit records the commercial terms of a specific assignment or engagement. Where it conflicts with "
        "the body of the Agreement on assignment-specific terms, this Exhibit controls for that assignment. An "
        "assignment record generated in the Company’s Operion platform and accepted by the Contractor serves the "
        "same purpose as a completed copy of this Exhibit."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(10)
    keep_with_next(p)

    fields = [
        ("Contractor name", ""),
        ("Role", "☐ Driver    ☐ Helper    ☐ Other: ______________________"),
        ("Service description", ""),
        ("Effective date", ""),
        ("Compensation basis", "☐ Per assignment    ☐ Per day    ☐ Per hour    ☐ Per route    ☐ Other: ____________"),
        ("Rate", "$ ______________  per  __________________"),
        ("Payment schedule", "Weekly, on Fridays, per Section 9 unless stated otherwise here: ______________"),
        ("Approved expenses", "☐ None    ☐ As listed: ______________________________________________"),
        ("Required equipment", ""),
        ("Vehicle requirements", "☐ Not applicable    ☐ As stated: _________________________________"),
        ("Insurance requirements", "☐ Per Section 14    ☐ Additional: _________________________________"),
        ("Special customer, route,\nor safety requirements", ""),
    ]

    # Free-text fields get a taller row so there is real room to write in them.
    free_text = {"Service description", "Effective date", "Required equipment",
                 "Special customer, route,\nor safety requirements", "Contractor name"}
    table = doc.add_table(rows=len(fields), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(fields):
        set_row_height(table.rows[i], 0.42 if label in free_text else 0.30)
        left, right = table.rows[i].cells
        left.width = Inches(2.0)
        right.width = Inches(4.5)
        shade(left, "F4F4F6")
        for cell, text, bold in ((left, label, True), (right, value, False)):
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for j, line in enumerate(text.split("\n")):
                run = p.add_run(("\n" if j else "") + line)
                run.bold = bold
                run.font.size = Pt(9.5)
        for cell in (left, right):
            cell_margins(cell)
            set_cell_border(
                cell,
                top={"val": "single", "sz": 4, "color": "CCCCCC"},
                bottom={"val": "single", "sz": 4, "color": "CCCCCC"},
                left={"val": "single", "sz": 4, "color": "CCCCCC"},
                right={"val": "single", "sz": 4, "color": "CCCCCC"},
            )

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    run = p.add_run("Accepted for this assignment:")
    run.bold = True
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    keep_with_next(p)
    keep_together(p)

    for label in (
        "Contractor signature:  ______________________________     Date:  ______________",
        f"{COMPANY_NAME} representative:  _____________________     Date:  ______________",
    ):
        sp = doc.add_paragraph()
        run = sp.add_run(label)
        run.font.size = Pt(10)
        sp.paragraph_format.space_after = Pt(12)
        keep_together(sp)


def convert_to_pdf(docx_path: Path) -> Path:
    result = subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
        capture_output=True, text=True, timeout=300,
    )
    pdf_path = docx_path.with_suffix(".pdf")
    if result.returncode != 0 or not pdf_path.exists():
        raise SystemExit(f"PDF conversion failed:\n{result.stdout}\n{result.stderr}")
    return pdf_path


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = OUT_DIR / f"{BASENAME}.docx"
    build_document().save(docx_path)
    pdf_path = convert_to_pdf(docx_path)
    print(f"DOCX  {docx_path}  ({docx_path.stat().st_size:,} bytes)")
    print(f"PDF   {pdf_path}  ({pdf_path.stat().st_size:,} bytes)")
    with pdf_path.open("rb") as fh:
        head = fh.read(5)
    print(f"PDF magic: {head!r}  valid={head == b'%PDF-'}")


if __name__ == "__main__":
    sys.exit(main())
